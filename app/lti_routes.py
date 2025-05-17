from flask import (
    request, jsonify, redirect, Blueprint, session,
    render_template, send_file, flash, url_for
)

from app.storage import (
    load_assignment_data,
    save_assignment_data,
    store_pending_feedback,
    load_pending_feedback,
    load_all_pending_feedback
)

import subprocess
import os

def convert_docx_to_html_with_styles(docx_path):
    html_path = docx_path.replace(".docx", ".html")
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "html",
            "--outdir", os.path.dirname(docx_path),
            docx_path
        ], check=True)

        with open(html_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print("❌ DOCX conversion failed:", e)
        return f"<p>Conversion error: {e}</p>"

import os
import json
import jwt
import uuid
from jwt.algorithms import RSAAlgorithm
from jwt.exceptions import InvalidTokenError
from requests_oauthlib import OAuth1Session
import requests
import mammoth
import re
import openai
from io import BytesIO
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text
from datetime import datetime, timedelta  
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from supabase import create_client

from app.utils.prompt_builder import build_grading_prompt 
from app.supabase_client import upload_to_supabase
from app.utils.zerogpt_api import check_ai_with_gpt
from app.models import Assignment

# 🛡️ Safe normalization function for assignment titles
def normalize_title(title):
    title = title.replace("–", "-").replace("—", "-")
    title = re.sub(r'\s+', ' ', title)
    title = title.replace("\u00A0", " ")
    title = title.lower().strip()
    return title

# ✅ Load environment and initialize Supabase client
load_dotenv()
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

FERPA_SAFE_MODE = os.getenv("FERPA_SAFE_MODE", "false").lower() == "true"

# ✅ Define the LTI blueprint
lti = Blueprint('lti', __name__)

def load_assignment_config(assignment_title):
    try:
        response = supabase.table("assignments").select("*").execute()
        if response.data:
            for assignment in response.data:
                stored_title = assignment.get("assignment_title", "")
                if normalize_title(stored_title) == normalize_title(assignment_title):
                    return assignment
    except Exception as e:
        print("❌ Error loading assignment config:", e)
    return None

ALLOWED_CLIENT_IDS = os.getenv("CLIENT_ID", "").split(",")

@lti.route("/login", methods=["POST"])
def login():
    print("🔐 /login route hit")
    issuer = request.form.get("iss")
    login_hint = request.form.get("login_hint")
    target_link_uri = request.form.get("target_link_uri")
    client_id = request.form.get("client_id")
    if client_id not in ALLOWED_CLIENT_IDS:
        return f"❌ Invalid client ID: {client_id}", 403
    
    lti_message_hint = request.form.get("lti_message_hint")

    if not all([issuer, login_hint, target_link_uri, client_id]):
        return "❌ Missing required LTI launch parameters", 400

    redirect_url = (
        f"{issuer}/mod/lti/auth.php?"
        f"scope=openid&response_type=id_token&client_id={client_id}&"
        f"redirect_uri={target_link_uri}&login_hint={login_hint}&state=state123&"
        f"response_mode=form_post&nonce=nonce123&prompt=none&"
        f"lti_message_hint={lti_message_hint}"
    )

    print(f"➡️ Redirecting to: {redirect_url}")
    return redirect(redirect_url)

@lti.before_app_request
def debug_session_state():
    print("🧪 SESSION CONTENTS:", dict(session))

def register_lti_routes(app):
    app.register_blueprint(lti)
    print("✅ LTI routes registered")

@lti.route("/.well-known/jwks.json", methods=["GET"])
def jwks():
    print("📡 Serving JWKS route...")
    with open("app/keys/jwks.json", "r") as f:
        jwks_data = json.load(f)
    return jsonify(jwks_data)

@lti.route("/.well-known/openid-configuration", methods=["GET"])
def openid_configuration():
    tool_url = os.environ["TOOL_URL"]
    return jsonify({
        "issuer": tool_url,
        "authorization_endpoint": f"{tool_url}/login",
        "token_endpoint": f"{tool_url}/token",
        "jwks_uri": f"{tool_url}/.well-known/jwks.json"
    })

@lti.route("/launch", methods=["POST"])
def launch():
    print("🚀 /launch hit")

    jwt_token = request.form.get("id_token")
    if not jwt_token:
        return "❌ Error: No id_token (JWT) received in launch request.", 400

    print("🌍 DEBUG - PLATFORM_ISS =", os.getenv("PLATFORM_ISS"))

    # Decode JWT without verifying (for debugging/platform info)
    try:
        unverified = jwt.decode(jwt_token, options={"verify_signature": False})
        print("🔍 UNVERIFIED JWT ISS:", unverified.get("iss"))
    except Exception as e:
        print("❌ Failed to decode unverified JWT:", str(e))
        return "❌ Unable to decode token.", 400

    # Fetch public key
    jwks_url = f"{os.getenv('PLATFORM_ISS')}/mod/lti/certs.php"
    try:
        jwks_response = requests.get(jwks_url)
        jwks_response.raise_for_status()
        jwks = jwks_response.json()
        kid = jwt.get_unverified_header(jwt_token).get("kid")
        public_key = next(
            (jwt.algorithms.RSAAlgorithm.from_jwk(json.dumps(k)) for k in jwks.get("keys", []) if k.get("kid") == kid),
            None
        )
        if not public_key:
            return "❌ No matching public key found in JWKS", 400
    except requests.RequestException as e:
        return f"❌ Could not fetch JWKS: {str(e)}", 400

    # Verify token signature and claims
    try:
        aud = jwt.decode(jwt_token, public_key, algorithms=["RS256"], options={"verify_aud": False}).get("aud")
        valid_client_ids = [cid.strip() for cid in os.getenv("CLIENT_IDS", "").split(",")]
        if aud not in valid_client_ids:
            return f"❌ Invalid client ID: {aud}", 403

        decoded = jwt.decode(jwt_token, public_key, algorithms=["RS256"], audience=aud, issuer=os.getenv("PLATFORM_ISS"))
        session["launch_data"] = decoded
        print("🧪 Saved launch_data keys to session:", list(decoded.keys()))

        print("✅ JWT verified")
        print(json.dumps(decoded, indent=2))
        roles = decoded.get("https://purl.imsglobal.org/spec/lti/claim/roles", [])
        is_instructor = any("Instructor" in role for role in roles)
        user_role = "instructor" if is_instructor else "student"
        session["tool_role"] = user_role
        print("🎓 LTI role detected:", user_role)


        # Send role to Supabase session so RLS works
        supabase.postgrest.rpc("set_config", {
            "key": "request.jwt.claims.role",
            "value": user_role
        })


    except Exception as e:
        return f"❌ Invalid JWT: {str(e)}", 400

    # ✅ Detect platform
    iss = decoded.get("iss", "")
    if "instructure.com" in iss:
        session["platform"] = "canvas"
    elif "moodle" in iss:
        session["platform"] = "moodle"
    else:
        session["platform"] = "unknown"

    # ✅ Store student info and AGS values
    session["student_id"] = decoded.get("sub")

    ags_claim = decoded.get("https://purl.imsglobal.org/spec/lti-ags/claim/endpoint", {})
    if ags_claim:
        session["lineitem_url"] = ags_claim.get("lineitem")
        print("🧠 AGS lineitem URL stored.")

    # Optional: capture token_url for future AGS access token
    custom_claims = decoded.get("https://purl.imsglobal.org/spec/lti/claim/custom", {})
    token_url = custom_claims.get("token_url")
    if token_url:
        session["token_url"] = token_url
        print("🔐 Token URL captured (for future AGS use).")

    # ✅ Load assignment config for launch UI
    assignment_title = decoded.get("https://purl.imsglobal.org/spec/lti/claim/resource_link", {}).get("title", "").strip()
    assignment_config = load_assignment_config(assignment_title)
    tinymce_api_key = os.getenv("TINYMCE_API_KEY")

    print("📦 Loaded config:", assignment_config)

    return render_template(
        "launch.html",
        activity_name=assignment_title,
        user_name=decoded.get("given_name", "Student"),
        user_roles=decoded.get("https://purl.imsglobal.org/spec/lti/claim/roles", []),
        assignment_config=assignment_config,
        tinymce_api_key=tinymce_api_key
    )


@lti.route("/student-test", methods=["GET"])
def student_test_upload():
    # Simulate a student launch with mock data
    session["tool_role"] = "student"
    user_role = "instructor"  # or "student"
    print("🎓 LTI role detected:", user_role)

    session["launch_data"] = {
        "https://purl.imsglobal.org/spec/lti/claim/resource_link": {
            "title": "Test Assignment"
        },
        "https://purl.imsglobal.org/spec/lti/claim/roles": ["Student"],
        "given_name": "Test User"
    }

    assignment_config = load_assignment_config("Test Assignment")

    return render_template(
        "launch.html",
        user_roles=["Student"],
        requires_persona=assignment_config.get("requires_persona", False) if assignment_config else False,
        assignment_config=assignment_config
    )


def get_total_points_from_rubric(rubric_json):
    """
    Calculates the total points from a rubric JSON file.
    Assumes each criterion has a 'points' field.
    """
    total_points = 0
    for criterion in rubric_json.get("criteria", []):
        points = criterion.get("points", 0)
        try:
            total_points += int(points)
        except (ValueError, TypeError):
            continue
    return total_points

@lti.route("/grade-docx", methods=["POST"])
def grade_docx():
    print(f"🔐 FERPA_SAFE_MODE: {FERPA_SAFE_MODE}")

    resource_link = session.get("launch_data", {}).get("https://purl.imsglobal.org/spec/lti/claim/resource_link", {})
    title_from_claim = resource_link.get("title")
    id_from_claim = resource_link.get("id")
    assignment_title = str(title_from_claim or f"Assignment-{id_from_claim}" or "Untitled Assignment").strip()

    assignment_title = str(assignment_title).strip()
    print(f"🧪 Final resolved assignment_title: {assignment_title}")

    print(f"🌐 RAW assignment_title FROM LAUNCH: {repr(assignment_title)}")
    assignment_title = normalize_title(assignment_title)
    assignment_config = load_assignment_config(assignment_title)

    if not assignment_config:
        return "❌ Assignment not found. Please contact your instructor.", 400

    delay_setting = assignment_config.get("delay_posting", "immediate")
    delay_map = {
        "immediate": 0,
        "1m": 0.0166,
        "12h": 12,
        "24h": 24,
        "36h": 36,
        "48h": 48
    }
    delay_hours = delay_map.get(delay_setting, 0)

    print("🧪 Grading assignment:", assignment_title)
    print("🧪 Assignment config loaded:", assignment_config)

    if not assignment_config.get("rubric_file", "").strip():
        return f"❌ Assignment setup incomplete. Missing configuration or rubric for: {assignment_title}", 400

    print("📅 /grade-docx hit")
    rubric_url = assignment_config.get("rubric_file", "")
    print("🧪 Rubric URL to download:", rubric_url)

    file = request.files.get("file")
    inline_text = request.form.get("inline_text", "").strip()
    reference_data = ""

    if not file and not inline_text:
        return "❌ Please submit either a file or inline response.", 400

    try:
        student_file_url = None  # initialize in case it's inline

        if file:
            filename = file.filename.lower()

            # Upload to Supabase
            from werkzeug.utils import secure_filename
            from io import BytesIO

            file_ext = os.path.splitext(filename)[-1]
            safe_name = secure_filename(filename)
            unique_path = f"submissions/{str(uuid.uuid4())}_{safe_name}"
            file_bytes = file.read()

            supabase.storage.from_("submissions").upload(unique_path, file_bytes)

            # Construct public URL (you can use env var or hardcode the project ID)
            SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID", "your-project-name.supabase.co")
            student_file_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/submissions/{unique_path}"

            # Extract text
            if file_ext == ".docx":
                doc = Document(BytesIO(file_bytes))
                full_text = "\n".join([para.text for para in doc.paragraphs])
            elif file_ext == ".pdf":
                full_text = extract_pdf_text(BytesIO(file_bytes))
            else:
                return "❌ Unsupported file type. Please upload .docx or .pdf", 400

        elif inline_text:
            full_text = inline_text


        ai_check_result = check_ai_with_gpt(full_text)
        print("🤖 AI Detection Result:", ai_check_result)

    except Exception as e:
        return f"❌ Failed to extract or process submission: {str(e)}", 500

    import tempfile, requests
    rubric_path = None
    if rubric_url:
        try:
            file_ext = rubric_url.split("?")[0].split(".")[-1]
            clean_filename = rubric_url.split("/")[-1].split("?")[0]
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{clean_filename}")
            response = requests.get(rubric_url)

            print("🧪 Rubric download HTTP status:", response.status_code)
            print("🧪 Rubric content size:", len(response.content))

            if response.status_code != 200 or len(response.content) == 0:
                return f"❌ Failed to download rubric from Supabase. Status {response.status_code}", 500

            temp_file.write(response.content)
            temp_file.close()
            rubric_path = temp_file.name

        except Exception as e:
            return f"❌ Exception while downloading rubric: {str(e)}", 500
    else:
        return "❌ No rubric file found for this assignment.", 400

    grading_difficulty = assignment_config.get("grading_difficulty", "balanced")
    student_level = assignment_config.get("student_level", "college")
    feedback_tone = assignment_config.get("feedback_tone", "supportive")
    ai_notes = assignment_config.get("ai_notes", "")

    try:
        if rubric_path.endswith(".json"):
            with open(rubric_path, "r") as f:
                rubric_json = json.load(f)
            rubric_text = "\n".join([f"- {c['description']}" for c in rubric_json.get("criteria", [])])
            rubric_total_points = get_total_points_from_rubric(rubric_json)
        elif rubric_path.endswith(".docx"):
            doc = Document(rubric_path)
            rubric_text = "\n".join([para.text for para in doc.paragraphs])
            rubric_total_points = assignment_config.get("total_points")
        elif rubric_path.endswith(".pdf"):
            rubric_text = extract_pdf_text(rubric_path)
            rubric_total_points = assignment_config.get("total_points")
        else:
            return "❌ No total points found. Please upload a .json rubric or specify a total in the dashboard.", 400

        if rubric_total_points is None:
            return "❌ This assignment does not have a total point value set. Please edit it in the dashboard.", 400
        rubric_total_points = int(str(rubric_total_points).strip())
        if rubric_total_points <= 0:
            raise ValueError
        print("✅ rubric_total_points confirmed as:", rubric_total_points)

    except Exception as e:
        return f"❌ Failed to load rubric file: {str(e)}", 500

    prompt = f"""
    You are a helpful AI grader.

    Assignment Title: {assignment_title}
    Grading Difficulty: {grading_difficulty}
    Student Level: {student_level}
    Feedback Tone: {feedback_tone}
    Total Points: {rubric_total_points}

    Rubric:
    {rubric_text}
    """
    if ai_notes:
        prompt += f"\nInstructor Notes:\n{ai_notes}\n"
    if reference_data:
        prompt += f"\nReference Scenario:\n{reference_data}\n"
    prompt += f"\nStudent Submission:\n---\n{full_text}\n---\n\nReturn your response in this format:\n\nScore: <number from 0 to {rubric_total_points}>\nFeedback: <detailed, encouraging, and helpful feedback>"

    try:
        openai.api_key = os.getenv("OPENAI_API_KEY")
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt.strip()}],
            temperature=0.5,
            max_tokens=1000
        )
        output = response["choices"][0]["message"]["content"]
        score_match = re.search(r"Score:\s*(\d{1,3})", output)
        score = int(score_match.group(1)) if score_match else 0
        feedback_match = re.search(r"Feedback:\s*(.+)", output, re.DOTALL)
        feedback = feedback_match.group(1).strip() if feedback_match else output.strip()
    except openai.error.OpenAIError as e:
        return f"❌ GPT error: {str(e)}", 500

    if not session.get("student_id"):
        fallback_id = session.get("launch_data", {}).get("sub")
        print("⚠️ session[\"student_id\"] missing. Using fallback:", fallback_id)
        session["student_id"] = fallback_id

    supabase.rpc("set_client_uid", {"uid": session["student_id"]}).execute()
    print("👤 Supabase client UID set to:", session["student_id"])

    submission_id = str(uuid.uuid4())
    submission_time = datetime.utcnow()
    release_time = submission_time + timedelta(hours=delay_hours)
    ready_to_post = delay_hours == 0 and not assignment_config.get("instructor_approval", False)
    pending = not ready_to_post

    submission_data = {
        "submission_id": submission_id,
        "student_id": session["student_id"],
        "assignment_title": assignment_title,
        "submission_time": submission_time.isoformat(),
        "score": score,
        "feedback": feedback,
        "submission_type": "inline" if inline_text else "file",
        "student_text": full_text,
        "student_file_url": student_file_url,
        "ai_check_result": None,
        "instructor_notes": "",
        "delay_hours": delay_hours,
        "ready_to_post": ready_to_post,
        "pending": pending,
        "reviewed": False,
        "release_time": release_time.isoformat()
    }

    print("🧪 INSERTING with student_id =", submission_data["student_id"])

    supabase.table("submissions").insert({
        "submission_id": submission_data["submission_id"],
        "student_id": submission_data["student_id"],
        "assignment_title": submission_data["assignment_title"],
        "submission_time": submission_data["submission_time"],
        "submission_type": submission_data["submission_type"],
        "delay_hours": delay_hours,
        "ready_to_post": delay_hours == 0 and not assignment_config.get("instructor_approval", False),
        "score": submission_data["score"],
        "feedback": submission_data["feedback"],
        "student_text": submission_data["student_text"],
        "student_file_url": student_file_url,
        "ai_check_result": submission_data["ai_check_result"],
        "instructor_notes": "",
        "pending": True,
        "reviewed": False,
        "release_time": (datetime.utcnow() + timedelta(hours=delay_hours)).isoformat()
    }).execute()

    log_gpt_interaction(assignment_title, prompt, feedback, score)

    if assignment_config.get("instructor_approval"):
        return render_template("feedback.html", pending_message="✅ This submission requires instructor review. Your feedback will be posted after approval.")
    elif delay_hours > 0:
        return render_template("feedback.html", pending_message=f"⏳ This submission will be released after {delay_hours} hour(s).")
    else:
        post_grade_to_lms(session, score, feedback)
        return render_template("feedback.html", score=score, feedback=feedback, rubric_total_points=rubric_total_points, user_roles=session.get("launch_data", {}).get("https://purl.imsglobal.org/spec/lti/claim/roles", []))





@lti.route("/review-feedback", methods=["GET", "POST"])
def review_feedback():
    session["tool_role"] = "instructor"  # TEMP: force instructor role for local/demo

    pending_path = os.path.join("rubrics", "pending_reviews.json")
    all_reviews = []
    if os.path.exists(pending_path):
        with open(pending_path, "r") as f:
            all_reviews = json.load(f)

    # Extract assignment titles from review queue
    assignment_titles = sorted(set(r["assignment_title"] for r in all_reviews))

    selected_title = request.form.get("assignment_title") if request.method == "POST" else None
    filtered_reviews = [r for r in all_reviews if r["assignment_title"] == selected_title] if selected_title else all_reviews

    if "review_index" not in session:
        session["review_index"] = 0

    if request.method == "POST" and "nav" in request.form:
        nav = request.form.get("nav")
        session["review_index"] = max(0, min(session["review_index"] + (1 if nav == "next" else -1), len(filtered_reviews) - 1))

    if request.method == "POST" and request.form.get("action") == "Approve and Post":
        student_id = request.form.get("student_id")
        assignment_title = request.form.get("assignment_title")
        allow_inline = request.form.get("allow_inline_submission", "no") == "yes"
        new_score = int(request.form.get("score"))
        new_feedback = request.form.get("feedback")

        matched = next((r for r in all_reviews if r["student_id"] == student_id and r["assignment_title"] == assignment_title), None)
        if matched:
            all_reviews.remove(matched)
            ags_claim = session.get("launch_data", {}).get("https://purl.imsglobal.org/spec/lti-ags/claim/endpoint")
            if ags_claim and "lineitem" in ags_claim:
                try:
                    lineitem_url = ags_claim["lineitem"].split("?")[0] + "/scores"
                    private_key_path = os.path.join("app", "keys", "private_key.pem")
                    with open(private_key_path, "r") as f:
                        private_key = f.read()
                    oauth = OAuth1Session(
                        client_key=os.getenv("CLIENT_ID"),
                        signature_method="RSA-SHA1",
                        rsa_key=private_key,
                        signature_type="auth_header"
                    )
                    score_payload = {
                        "userId": student_id,
                        "scoreGiven": new_score,
                        "scoreMaximum": assignment_config.get("total_points"),
                        "activityProgress": "Completed",
                        "gradingProgress": "FullyGraded",
                        "timestamp": datetime.utcnow().isoformat() + "Z"
                    }
                    oauth.post(
                        lineitem_url,
                        json=score_payload,
                        headers={"Content-Type": "application/vnd.ims.lis.v1.score+json"}
                    )
                except Exception as e:
                    print("❌ Grade post failed:", str(e))

        with open(pending_path, "w") as f:
            json.dump(all_reviews, f, indent=2)
        session["review_index"] = max(0, min(session["review_index"], len(filtered_reviews) - 1))

    current_review = filtered_reviews[session.get("review_index", 0)] if filtered_reviews else None

    return render_template(
        "review_feedback.html",
        assignment_titles=assignment_titles,
        selected_title=selected_title,
        current_review=current_review
    )

@lti.route("/post-grade", methods=["POST"])
def post_grade():
    score = int(request.form.get("score", 0))
    feedback = request.form.get("feedback", "")
    launch_data = session.get("launch_data", {})
    ags_claim = launch_data.get("https://purl.imsglobal.org/spec/lti-ags/claim/endpoint")

    if ags_claim and "lineitem" in ags_claim:
        try:
            lineitem_url = ags_claim["lineitem"].split("?")[0] + "/scores"
            private_key_path = os.path.join("app", "keys", "private_key.pem")
            with open(private_key_path, "r") as f:
                private_key = f.read()

            oauth = OAuth1Session(
                client_key=os.getenv("CLIENT_ID"),
                signature_method="RSA-SHA1",
                rsa_key=private_key,
                signature_type="auth_header"
            )

            score_payload = {
                "userId": launch_data.get("sub"),
                "scoreGiven": score,
                "scoreMaximum": assignment_config.get("total_points"),
                "activityProgress": "Completed",
                "gradingProgress": "FullyGraded",
                "timestamp": datetime.utcnow().isoformat() + "Z"
            }

            ags_response = oauth.post(
                lineitem_url,
                json=score_payload,
                headers={"Content-Type": "application/vnd.ims.lis.v1.score+json"}
            )
            ags_response.raise_for_status()
            print("✅ Manual grade posted by instructor.")
        except Exception as e:
            print("❌ Instructor grade post failed:", str(e))

    return render_template(
    "feedback.html",
    score=None,
    feedback=None,
    rubric_total_points=None,
    user_roles=session.get("launch_data", {}).get("https://purl.imsglobal.org/spec/lti/claim/roles", []),
    pending_message="✅ Your assignment was submitted successfully! Your grade and personalized feedback will be available after the review window."
)

@lti.route("/assignment-config", methods=["GET", "POST"])
def assignment_config():
    session["tool_role"] = "instructor"  # TEMP: for local testing
    tool_role = session.get("tool_role")
    if tool_role != "instructor":
        return "❌ Access denied. Instructors only.", 403

    rubric_index_path = os.path.join("rubrics", "rubric_index.json")
    rubric_folder = os.path.join("rubrics")

    if not os.path.exists(rubric_index_path):
        rubric_index = []
    else:
        with open(rubric_index_path, "r") as f:
            rubric_index = json.load(f)

    # Handle new assignment trigger
    if request.method == "GET" and request.args.get("new") == "1":
        rubric_index.insert(0, {
            "assignment_title": "",
            "rubric_file": "",
            "total_points": None,
            "instructor_approval": False,
            "requires_persona": False,
            "faith_integration": False,
            "grading_difficulty": "balanced",
            "student_level": "college",
            "feedback_tone": "supportive",
            "ai_notes": ""
        })
        with open(rubric_index_path, "w") as f:
            json.dump(rubric_index, f, indent=2)

    if request.method == "POST":
        # 📥 Get form fields
        assignment_title = request.form.get("assignment_title", "").strip()
        total_points_raw = request.form.get("total_points", "").strip()
        if not total_points_raw.isdigit():
            return "❌ Please enter a numeric total point value in the dashboard.", 400
        total_points = int(total_points_raw)
        instructor_approval = request.form.get("instructor_approval") == "on"
        requires_persona = request.form.get("requires_persona") == "on"
        faith_integration = request.form.get("faith_integration") == "on"
        grading_difficulty = request.form.get("grading_difficulty", "balanced")
        student_level = request.form.get("student_level", "college")
        feedback_tone = request.form.get("feedback_tone", "").strip()
        ai_notes = request.form.get("ai_notes", "").strip()
        allow_inline = request.form.get("allow_inline_submission", "no") == "yes"
        delay_posting = request.form.get("delay_posting", "immediate")


        rubric_file = request.files.get("rubric_file")
        saved_rubric_filename = None

        if rubric_file:
            safe_name = f"{assignment_title.lower().replace(' ', '_')}_{rubric_file.filename}"
            rubric_path = os.path.join(rubric_folder, safe_name)
            rubric_file.save(rubric_path)
            saved_rubric_filename = safe_name

        updated = False
        for entry in rubric_index:
            if entry["assignment_title"].lower() == assignment_title.lower():
                entry.update({
                    "rubric_file": saved_rubric_filename,
                    "total_points": total_points,
                    "instructor_approval": instructor_approval,
                    "requires_persona": requires_persona,
                    "faith_integration": faith_integration,
                    "grading_difficulty": grading_difficulty,
                    "student_level": student_level,
                    "feedback_tone": feedback_tone,
                    "ai_notes": ai_notes,
                    "delay_posting": delay_posting,
                    "allow_inline_submission": allow_inline,

                })
                updated = True
                break

        if not updated:
            rubric_index.append({
                "assignment_title": assignment_title,
                "rubric_file": saved_rubric_filename,
                "total_points": total_points,
                "instructor_approval": instructor_approval,
                "requires_persona": requires_persona,
                "faith_integration": faith_integration,
                "grading_difficulty": grading_difficulty,
                "student_level": student_level,
                "feedback_tone": feedback_tone,
                "ai_notes": ai_notes,
                "delay_posting": delay_posting,
                "allow_inline_submission": allow_inline,

            })

        with open(rubric_index_path, "w") as f:
            json.dump(rubric_index, f, indent=2)

    return render_template("assignment_config.html", rubric_index=rubric_index)

@lti.route("/test-grader", methods=["GET", "POST"])
def test_grader():
    rubric_index_path = os.path.join("rubrics", "rubric_index.json")
    rubric_index = []

    if os.path.exists(rubric_index_path):
        with open(rubric_index_path, "r") as f:
            rubric_index = json.load(f)

    selected_config = None
    gpt_prompt = ""
    gpt_feedback = ""
    gpt_score = None

    if request.method == "POST":
        assignment_title = request.form.get("assignment_title")
        submission_text = request.form.get("submission_text", "").strip()

        selected_config = next((cfg for cfg in rubric_index if cfg["assignment_title"] == assignment_title), None)

        if not selected_config:
            return "❌ No config found for that assignment.", 400
        
        # ✅ Check if grading should be delayed
        # ✅ Check if grading should be delayed
        delay_setting = selected_config.get("delay_posting", "immediate")

        delay_map = {
            "immediate": 0,
            "12h": 12,
            "24h": 24,
            "36h": 36,
            "48h": 48
        }

        delay_hours = delay_map.get(delay_setting, 0)

        if delay_hours > 0:
            from datetime import datetime, timedelta
            from app.storage import store_pending_feedback  # Ensure this import is valid

            release_time = datetime.utcnow() + timedelta(hours=delay_hours)

            store_pending_feedback(
                assignment_title=assignment_title,
                student_id="test_user",  # Replace with real student ID later
                feedback="(To be generated at release time)",
                score=None,
                release_time=release_time.isoformat()
            )

            gpt_feedback = f"⏳ Feedback for '{assignment_title}' will be generated after {delay_hours} hour(s)."
            gpt_score = None

            return render_template(
                "test_grader.html",
                rubric_index=rubric_index,
                selected_config=selected_config,
                gpt_prompt="",
                gpt_feedback=gpt_feedback,
                gpt_score=gpt_score
            )

        # Load rubric
        rubric_text = ""
        rubric_path = os.path.join("rubrics", selected_config.get("rubric_file", ""))
        try:
            if rubric_path.endswith(".json"):
                with open(rubric_path, "r") as f:
                    rubric_json = json.load(f)
                rubric_text = "\n".join([f"- {c['description']}" for c in rubric_json.get("criteria", [])])
            elif rubric_path.endswith(".docx"):
                doc = Document(rubric_path)
                rubric_text = "\n".join([para.text for para in doc.paragraphs])
            elif rubric_path.endswith(".pdf"):
                rubric_text = extract_pdf_text(rubric_path)
        except:
            rubric_text = "(Unable to load rubric.)"

        # Build GPT prompt
        prompt = f"""
You are a helpful AI grader.

Assignment Title: {assignment_title}
Grading Difficulty: {selected_config.get("grading_difficulty")}
Student Level: {selected_config.get("student_level")}
Feedback Tone: {selected_config.get("feedback_tone")}
Total Points: {selected_config.get("total_points")}

Rubric:
{rubric_text}
"""

        if selected_config.get("ai_notes"):
            prompt += f"""

Instructor Notes:
{selected_config['ai_notes']}
"""

        prompt += f"""

Student Submission:
---
The following is the student's full submission. Please preserve paragraph formatting, line breaks, and indentation when analyzing or quoting their writing.

\"\"\"{submission_text}\"\"\"
---

Return your response in this format:

Score: <number from 0 to {selected_config.get("total_points")}>
Feedback: <detailed, helpful feedback in paragraph form>

(Note: A future version of this tool may request table-based feedback. If no such instruction is present, return standard narrative feedback only.)
"""

        gpt_prompt = prompt.strip()

        try:
            openai.api_key = os.getenv("OPENAI_API_KEY")
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": gpt_prompt}],
                temperature=0.5,
                max_tokens=1000
            )
            output = response["choices"][0]["message"]["content"]

            score_match = re.search(r"Score:\s*(\d{1,3})", output)
            gpt_score = int(score_match.group(1)) if score_match else None

            feedback_match = re.search(r"Feedback:\s*(.+)", output, re.DOTALL)
            gpt_feedback = feedback_match.group(1).strip() if feedback_match else output.strip()

            log_gpt_interaction(assignment_title, gpt_prompt, gpt_feedback, gpt_score)

        except Exception as e:
            gpt_feedback = f"❌ GPT error: {str(e)}"

    return render_template(
        "test_grader.html",
        rubric_index=rubric_index,
        selected_config=selected_config,
        gpt_prompt=gpt_prompt,
        gpt_feedback=gpt_feedback,
        gpt_score=gpt_score if 'gpt_score' in locals() else None
    )

@lti.route("/save-assignment", methods=["POST"])
def save_assignment():
    assignment_title = request.form.get("assignment_title", "").strip()
    if not assignment_title:
        return "❌ Assignment title is required", 400

    grading_difficulty = request.form.get("grading_difficulty")
    grade_level = request.form.get("grade_level")
    total_points = int(request.form.get("total_points", "0"))
    instructor_approval = request.form.get("instructor_approval") == "true"
    gospel_enabled = request.form.get("gospel_enabled") == "true"
    custom_ai = request.form.get("custom_ai", "")

    rubric_file = request.files.get("rubric_upload")
    additional_file = request.files.get("additional_files")

    allow_inline = request.form.get("allow_inline_submission", "no") == "yes"


    from werkzeug.utils import secure_filename
    upload_dir = os.path.join("uploads", secure_filename(assignment_title))
    os.makedirs(upload_dir, exist_ok=True)

    rubric_url = ""
    additional_url = ""

    # Handle rubric upload
    if rubric_file and rubric_file.filename:
        rubric_filename = secure_filename(rubric_file.filename)
        rubric_filename = rubric_filename.replace(" ", "_")
        rubric_path = os.path.join(upload_dir, rubric_filename)
        rubric_file.save(rubric_path)
        rubric_url = upload_to_supabase(rubric_path, rubric_filename)
        if rubric_url:
            rubric_url = rubric_url.rstrip("?")
            if "rubrics/rubrics/" in rubric_url:
                rubric_url = rubric_url.replace("rubrics/rubrics/", "rubrics/")

    if not rubric_url:
        flash("❌ Rubric upload failed. Please try uploading the rubric again.", "error")
        return redirect(url_for('lti.view_assignments'))

    # Handle additional file upload
    if additional_file and additional_file.filename:
        additional_filename = secure_filename(additional_file.filename)
        additional_filename = additional_filename.replace(" ", "_")
        additional_path = os.path.join(upload_dir, additional_filename)
        additional_file.save(additional_path)
        additional_url = upload_to_supabase(additional_path, additional_filename)
        if additional_url:
            additional_url = additional_url.rstrip("?")
            if "attachments/attachments/" in additional_url:
                additional_url = additional_url.replace("attachments/attachments/", "attachments/")

    rubric_url = rubric_url or ""
    additional_url = additional_url or ""

    # ✅ Save directly to Supabase and capture response
    response = supabase.table("assignments").insert({
        
        "assignment_title": assignment_title,
        "rubric_file": rubric_url,
        "additional_file": additional_url,
        "total_points": total_points,
        "instructor_approval": instructor_approval,
        "requires_persona": False,
        "faith_integration": gospel_enabled,
        "grading_difficulty": grading_difficulty,
        "student_level": grade_level,
        "feedback_tone": "supportive",
        "ai_notes": custom_ai,
        "allow_inline_submission": allow_inline 
    }).execute()
    

    # ✅ Debug print statements BEFORE redirect
    print("🧪 Saving assignment:", assignment_title)
    print("🧪 Rubric URL:", rubric_url)
    print("🧪 Additional file URL:", additional_url)

    # ✅ Check if insert succeeded
    if hasattr(response, 'error') and response.error:
        flash(f"❌ Error saving assignment: {response.error['message']}", "error")
        return redirect(url_for('lti.view_assignments'))

    # ✅ No error, proceed normally
    flash("✅ Assignment saved successfully.", "success")
    return redirect(f"/admin-dashboard?success={assignment_title}")


@lti.route("/admin-dashboard", methods=["GET", "POST"])
def admin_dashboard():
    session["tool_role"] = "instructor"  # TEMP for local testing

    if session.get("tool_role") != "instructor":
        return "❌ Access denied. Instructors only.", 403

    try:
        # ✅ NEW: Fetch assignments live from Supabase
        response = supabase.table("assignments").select("*").execute()
        if response.data is None:
            rubric_index = []
        else:
            rubric_index = response.data
    except Exception as e:
        print("❌ Supabase fetch error (assignments):", e)
        flash("❌ Error loading assignments.", "danger")
        rubric_index = []

    try:
        # ✅ Fetch pending feedback
        response = supabase.table("submissions").select("*").eq("pending", True).execute()
        pending_feedback = response.data or []
    except Exception as e:
        print("❌ Supabase fetch error (pending submissions):", e)
        flash("❌ Error loading pending submissions.", "danger")
        pending_feedback = []


    # ✅ Load activity logs
    log_path = os.path.join(os.path.dirname(__file__), "..", "logs", "activity_log.json")
    if os.path.exists(log_path):
        with open(log_path, "r") as f:
            activity_logs = json.load(f)
    else:
        activity_logs = []

    pending_count = len(pending_feedback)
    approved_count = sum(1 for r in rubric_index if r.get("instructor_approval"))

    return render_template("admin_dashboard.html",
                           rubric_index=rubric_index,
                           pending_feedback=pending_feedback,
                           pending_count=pending_count,
                           approved_count=approved_count,
                           activity_logs=activity_logs)


@lti.route("/instructor-review", methods=["GET", "POST"])
def instructor_review():
    response = supabase.table("submissions").select("*").eq("pending", True).execute()
    reviews = response.data or []
    submission_id = request.args.get("submission_id")
    current_review = None
    next_id = None

    # 🔁 Recalculate reviews fresh from Supabase after any Accept
    if submission_id:
        # Find the index of the current submission
        for i, review in enumerate(reviews):
            if review["submission_id"] == submission_id:
                current_review = review
                if i + 1 < len(reviews):
                    next_id = reviews[i + 1]["submission_id"]
                break

    # Default to first if no match
    if not current_review and reviews:
        current_review = reviews[0]
        if len(reviews) > 1:
            next_id = reviews[1]["submission_id"]

    print("🧪 FINAL current_review =", current_review)
    print("🧪 Remaining reviews:", len(reviews))


    print(f"🧪 Number of pending reviews found: {len(reviews)}")

    submission_id = request.args.get("submission_id")

    if request.method == "POST":
        submission_id = request.form.get("submission_id")
        updated_score = int(request.form.get("score"))
        updated_feedback = request.form.get("feedback")

        supabase.table("submissions").update({
            "score": updated_score,
            "feedback": updated_feedback,
            "submission_time": datetime.utcnow().isoformat(),
            "reviewed": True,
            "pending": False
        }).eq("submission_id", submission_id).execute()

        supabase.table("submissions").update({
            "score": updated_score,
            "feedback": updated_feedback,
            "timestamp": datetime.utcnow().isoformat()
        }).eq("submission_id", submission_id).execute()

        return redirect(url_for("lti.instructor_review"))

    current_review = None
    if submission_id:
        current_review = next((r for r in reviews if r["submission_id"] == submission_id), None)
    elif reviews:
        current_review = reviews[0]
    
    print("🧪 current_review =", current_review)

    docx_html = None
    pdf_url = None
    docx_pages = []

    if current_review:
        file_url = str(current_review.get("student_file_url") or "")
        if file_url.endswith(".pdf"):
            pdf_url = file_url
        elif file_url.endswith(".docx"):
            try:
                response = requests.get(file_url)
                docx_path = "/tmp/temp.docx"
                with open(docx_path, "wb") as f:
                    f.write(response.content)

                full_html = convert_docx_to_html_with_styles(docx_path)

                # Split pages if LibreOffice used <div style="page-break-before: always">
                docx_pages = full_html.split('<div style="page-break-before: always"')

                # Re-add the opening div to each split page (except the first)
                for i in range(1, len(docx_pages)):
                    docx_pages[i] = '<div style="page-break-before: always"' + docx_pages[i]

                    print("📄 Total DOCX pages found:", len(docx_pages))
                # Remove the last page if it's empty
            
            except Exception as e:
                print("❌ DOCX rendering error:", e)


    return render_template(
    "instructor_review.html",
    current_review=current_review,
    reviews=reviews,
    docx_pages=docx_pages,
    pdf_url=pdf_url,
    next_id=next_id
)




@lti.route("/instructor-review/save-notes", methods=["POST"])
def instructor_save_notes():
    submission_id = request.form.get("submission_id")
    new_notes = request.form.get("notes", "")

    if not submission_id:
        return "❌ Submission ID missing", 400

    # Save notes to Supabase
    response = supabase.table("submissions").update({
        "instructor_notes": new_notes
    }).eq("submission_id", submission_id).execute()

    if hasattr(response, 'error') and response.error:
        return f"❌ Supabase error: {response.error.message}", 500

    # ✅ Optionally log activity
    from app.utils.logging import log_activity
    log_activity("Saved instructor notes", user="instructor", details=submission_id)

    return redirect("/admin-dashboard")



@lti.route("/export-configs", methods=["GET"])
def export_configs():
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        rubric_path = os.path.join(base_dir, "..", "rubrics", "rubric_index.json")

        print("🧾 Attempting to export:", rubric_path)

        if not os.path.exists(rubric_path):
            return "No configuration file found.", 404

        return send_file(
            rubric_path,
            mimetype="application/json",
            as_attachment=True,
            download_name="rubiqs_assignment_configs.json"
        )
    except Exception as e:
        print("❌ EXPORT ERROR:", str(e))
        return f"Export failed: {str(e)}", 500
    
@lti.route("/update-config", methods=["POST"])
def update_config():
    rubric_index_path = os.path.join("rubrics", "rubric_index.json")
    assignment_title = request.form.get("assignment_title")
    total_points_raw = request.form.get("total_points", "").strip()
    if not total_points_raw.isdigit():
        return "❌ Please enter a valid number for total points.", 400
    total_points = int(total_points_raw)

    ai_notes = request.form.get("ai_notes", "").strip()
    grading_difficulty = request.form.get("grading_difficulty", "balanced")
    student_level = request.form.get("student_level", "college")
    faith_raw = request.form.get("faith_integration", "false")
    faith_integration = True if faith_raw.lower() == "true" else False
    allow_inline = request.form.get("allow_inline_submission", "no") == "yes"
    delay_posting = request.form.get("delay_posting", "immediate")



    if not os.path.exists(rubric_index_path):
        return "Config file missing", 404

    with open(rubric_index_path, "r") as f:
        rubric_index = json.load(f)

    for entry in rubric_index:
        if entry["assignment_title"] == assignment_title:
            entry["total_points"] = total_points
            entry["ai_notes"] = ai_notes
            entry["grading_difficulty"] = grading_difficulty
            entry["student_level"] = student_level
            entry["faith_integration"] = faith_integration
            entry["allow_inline_submission"] = allow_inline

    with open(rubric_index_path, "w") as f:
        json.dump(rubric_index, f, indent=2)

    print("✅ Updated config for:", assignment_title)
    print("✅ allow_inline =", allow_inline)



    return redirect("/admin-dashboard")

def log_gpt_interaction(assignment_title, prompt, feedback, score=None):
    log_path = os.path.join("logs", "prompt_logs.json")
    os.makedirs("logs", exist_ok=True)

    log_entry = {
        "assignment_title": assignment_title,
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "prompt": prompt,
        "feedback": feedback,
        "score": score
    }

    logs = []
    if os.path.exists(log_path):
        with open(log_path, "r") as f:
            try:
                logs = json.load(f)
            except json.JSONDecodeError:
                logs = []

    logs.append(log_entry)

    with open(log_path, "w") as f:
        json.dump(logs, f, indent=2)


@lti.route("/test-store")
def test_store():
    import uuid
    from datetime import datetime
    from app.storage import store_pending_feedback


    submission_id = str(uuid.uuid4())
    test_data = {
        "submission_id": submission_id,
        "student_id": "tester001",
        "assignment_title": "Test Assignment",
        "timestamp": datetime.utcnow().isoformat(),
        "score": 88,
        "feedback": "Nice test work!",
        "student_text": "This is a test submission.",
        "ai_check_result": None
    }

    store_pending_feedback(submission_id, test_data)
    return f"✅ Submission saved: {submission_id}"

# In app/lti_routes.py
from flask import request, jsonify
from app.utils.zerogpt_api import check_ai_with_gpt


@lti.route("/scan-ai", methods=["POST"])
def scan_ai():
    data = request.get_json()
    text = data.get("text", "")

    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        import openai
        import os
        openai.api_key = os.getenv("OPENAI_API_KEY")

        prompt = f"""
You are an expert in AI-authorship detection. Your task is to analyze the following student submission and determine the likelihood that it was written using an AI tool (such as ChatGPT, Claude, or Gemini).

Be on the lookout for the following signs of AI-generated writing:
- Overly polished and generic phrasing
- Repetitive sentence structures or overly formal transitions
- Lack of personal anecdotes, emotional reflection, or specific experience
- Idealized writing with no personal struggle or ambiguity
- Balanced arguments without emotional conviction or imperfection
- Fluency that feels too even, with little variation in sentence length

Please respond with the following structure:
1. SCORE: Estimate a percentage likelihood from 0% (definitely human) to 100% (definitely AI-generated).
2. REASONING: Explain why you assigned this score.
3. EVIDENCE: Quote specific portions of the student's writing that influenced your judgment.

Text to analyze:
\"\"\"
{text}
\"\"\"
"""

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )

        result = response["choices"][0]["message"]["content"]
        return jsonify({"result": result})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@lti.route("/instructor-review-button", methods=["GET", "POST"])
def instructor_review_button():
    response = supabase.table("submissions").select("*").eq("pending", True).execute()
    reviews = response.data or []

    print(f"🧪 Number of pending reviews found: {len(reviews)}")  # ✅ Add this

    submission_id = request.args.get("submission_id")

    if request.method == "POST":
        submission_id = request.form.get("submission_id")
        updated_score = int(request.form.get("score"))
        updated_feedback = request.form.get("feedback")

        supabase.table("submissions").update({
            "score": updated_score,
            "feedback": updated_feedback,
            "timestamp": datetime.utcnow().isoformat()
        }).eq("submission_id", submission_id).execute()

        return redirect(url_for("lti.instructor_review_button"))

    current_review = None
    if submission_id:
        current_review = next((r for r in reviews if r["submission_id"] == submission_id), None)
    elif reviews:
        current_review = reviews[0]

        return render_template("instructor_review.html", current_review=current_review, reviews=reviews)

def post_grade_to_lms(session, score, feedback):
    try:
        launch_data = session.get("launch_data", {})
        ags_claim = launch_data.get("https://purl.imsglobal.org/spec/lti-ags/claim/endpoint")

        if not ags_claim or "lineitem" not in ags_claim:
            print("⚠️ AGS info missing — cannot post grade.")
            return

        lineitem_url = ags_claim["lineitem"].split("?")[0] + "/scores"
        private_key_path = os.path.join("app", "keys", "private_key.pem")
        with open(private_key_path, "r") as f:
            private_key = f.read()

        oauth = OAuth1Session(
            client_key=os.getenv("CLIENT_ID"),
            signature_method="RSA-SHA1",
            rsa_key=private_key,
            signature_type="auth_header"
        )

        score_payload = {
            "userId": launch_data.get("sub"),
            "scoreGiven": score,
            "scoreMaximum": 100,  # Optional: You can customize this later
            "activityProgress": "Completed",
            "gradingProgress": "FullyGraded",
            "timestamp": datetime.utcnow().isoformat() + "Z"
        }

        response = oauth.post(
            lineitem_url,
            json=score_payload,
            headers={"Content-Type": "application/vnd.ims.lis.v1.score+json"}
        )

        if response.status_code >= 200 and response.status_code < 300:
            print("✅ Grade posted to LMS.")
        else:
            print("⚠️ AGS post failed:", response.text)

    except Exception as e:
        print("❌ Error in post_grade_to_lms():", str(e))


from app.database import SessionLocal
from app.models import Assignment

@lti.route("/edit-assignment/<int:assignment_id>", methods=["GET", "POST"])
def edit_assignment(assignment_id):
    if request.method == "POST":
        print("🚀 Save Assignment POST route hit")
        try:
            print("📥 POST data:", request.form)

            total_points = request.form.get("total_points", type=int)
            ai_notes = request.form.get("ai_notes", "")
            student_level = request.form.get("student_level")
            grading_difficulty = request.form.get("grading_difficulty")
            delay_posting = request.form.get("delay_posting", "immediate")
            allow_inline = request.form.get("allow_inline_submission", "no") == "yes"

            faith_raw = request.form.get("faith_integration", "false")
            faith_integration = True if faith_raw.lower() == "true" else False

            delay_raw = request.form.get("delay_posting", "immediate")

            # ✅ Update Supabase assignment record by ID
            response = supabase.table("assignments").update({
                "total_points": total_points,
                "ai_notes": ai_notes,
                "student_level": student_level,
                "grading_difficulty": grading_difficulty,
                "faith_integration": faith_integration,
                "delay_posting": delay_raw,
                "allow_inline_submission": allow_inline
            }).eq("id", assignment_id).execute()

            if hasattr(response, "error") and response.error:
                print("❌ Supabase error:", response.error.message)
                return f"❌ Supabase update error: {response.error.message}", 500

            print("✅ Assignment updated successfully")
            return redirect(url_for("lti.view_assignments"))

        except Exception as e:
            print("❌ Exception in edit_assignment:", e)
            return "Internal Server Error", 500

    # GET request: fetch assignment by ID
    response = supabase.table("assignments").select("*").eq("id", assignment_id).single().execute()
    if response.data is None:
        return "Assignment not found", 404

    assignment = response.data
    return render_template("edit_assignment.html", assignment=assignment)



@lti.route("/view-assignments")
def view_assignments():
    try:
        response = supabase.table("assignments").select("*").execute()
        assignments = response.data or []
    except Exception as e:
        print("❌ Supabase fetch error (view-assignments):", e)
        flash("❌ Error loading assignments.", "danger")
        assignments = []

    return render_template("view_assignments.html", assignments=assignments)


@lti.route('/delete-file', methods=['POST'])
def delete_file():
    import os
    from flask import request, jsonify
    from supabase import create_client
    from app.database import SessionLocal
    from app.models import Assignment

    try:
        data = request.get_json()
        print("📥 Incoming delete request:", data)

        assignment_id = data.get("assignment_id")
        file_type = data.get("file_type")

        if not assignment_id or file_type not in ["rubric", "additional"]:
            print("❌ Missing or invalid data")
            return jsonify({"success": False, "error": "Missing or invalid data"}), 400

        session = SessionLocal()
        assignment = session.query(Assignment).filter_by(id=assignment_id).first()

        if not assignment:
            print("❌ Assignment not found")
            return jsonify({"success": False, "error": "Assignment not found"}), 404

        supabase = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_KEY"))

        deleted_filename = ""
        if file_type == "rubric" and assignment.rubric_file:
            deleted_filename = assignment.rubric_file.split("/")[-1]
            assignment.rubric_file = None
        elif file_type == "additional" and assignment.additional_file:
            deleted_filename = assignment.additional_file.split("/")[-1]
            assignment.additional_file = None

        if deleted_filename:
            deleted_filename = deleted_filename.lstrip("/")
            print(f"🗑️ Attempting to delete file: {deleted_filename}")
            res = supabase.storage.from_("rubrics").remove([deleted_filename])
            print("✅ Supabase response:", res)

        session.commit()
        return jsonify({"success": True})

    except Exception as e:
        print("❌ Exception caught in /delete-file route:", e)
        return jsonify({"success": False, "error": str(e)}), 500

    finally:
        try:
            session.close()
        except:
            pass


@lti.route("/dev/add-notes-column")
def add_notes_column():
    try:
        supabase.rpc("alter_table_add_column", {
            "table_name": "submissions",
            "column_name": "instructor_notes",
            "column_type": "text"
        }).execute()
        return "✅ instructor_notes column added (or already exists)"
    except Exception as e:
        return f"❌ Error: {str(e)}"

@lti.route('/delete-assignment', methods=['POST'])
def delete_assignment():
    data = request.get_json()
    assignment_title = data.get('assignment_title')

    if not assignment_title:
        return jsonify({'success': False, 'error': 'Missing assignment title'}), 400

    try:
        response = supabase.table("assignments").delete().eq("assignment_title", assignment_title).execute()

        if hasattr(response, "error") and response.error:
            return jsonify({"success": False, "error": response.error.message}), 500

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500




@lti.route("/release-pending", methods=["GET"])
def release_pending_feedback():
    print("🚀 /release-pending triggered")

    try:
        from datetime import datetime
        now = datetime.utcnow().isoformat()

        response = supabase.table("submissions") \
            .select("*") \
            .eq("pending", True) \
            .lt("release_time", now) \
            .execute()

        if hasattr(response, "error") and response.error:
            print("❌ Supabase query error:", response.error.message)
            return f"Supabase query failed: {response.error.message}", 500

        pending = response.data or []
        print(f"📬 Found {len(pending)} entries eligible for release")

        released = 0

        for entry in pending:
            assignment_id = entry.get("assignment_id")
            student_id = entry.get("student_id")
            score = entry.get("score")
            feedback = entry.get("feedback")
            submission_id = entry.get("submission_id")

            if not all([assignment_id, student_id, score, feedback, submission_id]):
                print(f"⚠️ Skipping incomplete submission: {submission_id}")
                continue

            update_response = supabase.table("submissions").update({
                "pending": False,
                "reviewed": True,
                "released_at": now
            }).eq("submission_id", submission_id).execute()

            if hasattr(update_response, "error") and update_response.error:
                print(f"❌ Error updating submission {submission_id}: {update_response.error.message}")
                continue

            released += 1

        return f"✅ Released {released} submissions", 200

    except Exception as e:
        print("❌ Fatal error in release process:", str(e))
        return f"❌ Internal error: {str(e)}", 500

@lti.route("/run-delay-checker")
def run_delay_checker():
    from datetime import datetime, timedelta

    # Fetch submissions that are still waiting
    response = supabase.table("submissions").select("*").eq("ready_to_post", False).execute()

    if not response.data:
        print("✅ No submissions pending release.")
        return "✅ No pending submissions to check.", 200

    now = datetime.utcnow()
    updates_made = 0

    for submission in response.data:
        try:
            # Get the release_time from the database (this is already stored when you save submissions)
            release_time = datetime.fromisoformat(submission["release_time"].replace("Z", ""))  # Remove 'Z' if present

            # Check if the current time is greater than or equal to the release time
            if now >= release_time:
                # Update ready_to_post to True if delay has expired
                supabase.table("submissions").update({"ready_to_post": True}).eq("submission_id", submission["submission_id"]).execute()
                updates_made += 1

        except Exception as e:
            print(f"❌ Error checking submission {submission.get('submission_id')}: {str(e)}")

    print(f"✅ Delay check complete. Updated {updates_made} submissions.")

    return f"✅ Delay check complete. Updated {updates_made} submissions.", 200

import csv
from flask import make_response

@lti.route("/download-activity-log")
def download_activity_log():
    try:
        # Fetch all submissions
        response = supabase.table("submissions").select("*").order("timestamp", desc=True).execute()
        submissions = response.data or []

        # Build CSV content
        si = []
        headers = ["timestamp", "student_id", "assignment_title", "score", "reviewed", "pending"]

        # Add header row
        si.append(",".join(headers))

        for submission in submissions:
            row = [
                submission.get("timestamp", ""),
                submission.get("student_id", ""),
                submission.get("assignment_title", ""),
                str(submission.get("score", "")),
                str(submission.get("reviewed", "")),
                str(submission.get("pending", ""))
            ]
            si.append(",".join(row))

        csv_content = "\n".join(si)

        # Return CSV file
        response = make_response(csv_content)
        response.headers["Content-Disposition"] = "attachment; filename=activity_log.csv"
        response.headers["Content-type"] = "text/csv"
        return response

    except Exception as e:
        print("❌ Error generating activity log download:", e)
        return "❌ Failed to generate report.", 500

# 🧪 Triggering redeploy

@lti.route('/delete-submission', methods=['POST'])
def delete_submission():
    submission_id = request.form.get("submission_id")
    print("🧪 Received delete request for:", submission_id)

    if not submission_id:
        print("❌ Missing submission ID")
        return jsonify({"success": False, "error": "Missing submission ID"}), 400

    try:
        response = supabase.table("submissions")\
            .delete()\
            .eq("submission_id", submission_id)\
            .execute()

        print("🧪 Supabase delete response:", response)

        if hasattr(response, "error") and response.error:
            print("❌ Supabase delete error:", response.error.message)
            return jsonify({"success": False, "error": response.error.message}), 500

        print("✅ Deleted submission_id:", submission_id)
        return jsonify({"success": True}), 200

    except Exception as e:
        print("❌ Exception caught during deletion:", str(e))
        return jsonify({"success": False, "error": "Internal server error"}), 500


@lti.route("/test-insert")
def test_insert():
    session["student_id"] = "2"  # Replace with your real student ID
    supabase.rpc("set_client_uid", {"uid": session["student_id"]}).execute()

    import uuid
    submission_id = str(uuid.uuid4())

    test_submission = {
        "submission_id": submission_id,
        "student_id": session["student_id"],
        "assignment_title": "Test Assignment",
        "submission_time": datetime.utcnow().isoformat(),
        "submission_type": "inline",
        "delay_hours": 0,
        "ready_to_post": True,
        "score": 100,
        "feedback": "Great job!",
        "student_text": "Test submission",
        "ai_check_result": None,
        "instructor_notes": "",
        "pending": False,
        "reviewed": True,
        "release_time": datetime.utcnow().isoformat()
    }

    supabase.table("submissions").insert(test_submission).execute()
    return "✅ Insert succeeded"

@lti.route('/instructor-review/accept', methods=['POST'])
def accept_review():
    submission_id = request.form.get("submission_id", "").strip()
    print("🧪 Accepting submission_id:", submission_id)

    if not submission_id:
        return jsonify({"success": False, "error": "Missing submission ID"}), 400

    try:
        response = supabase.table("submissions")\
            .update({
                "pending": False,
                "reviewed": True
            })\
            .eq("submission_id", submission_id)\
            .execute()

        print("✅ Submission accepted and updated.")
        return jsonify({"success": True})

    except Exception as e:
        print("❌ Accept error:", str(e))
        return jsonify({"success": False, "error": "Server error"}), 500

